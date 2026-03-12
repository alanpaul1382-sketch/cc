#!/usr/bin/env python3
"""
Summarize a YouTube video and generate a .docx report with relevant images.

Usage:
    python summarize_video.py [VIDEO_URL]

If no URL is provided, defaults to: https://www.youtube.com/watch?v=SvXEAj7Mb48
"""

import argparse
import os
import re
import sys
import textwrap
from io import BytesIO
from urllib.parse import parse_qs, urlparse
from urllib.request import urlopen, Request

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from youtube_transcript_api import YouTubeTranscriptApi
import yt_dlp

DEFAULT_VIDEO_URL = "https://www.youtube.com/watch?v=SvXEAj7Mb48"
OUTPUT_DIR = "output"


def extract_video_id(url):
    """Extract the video ID from a YouTube URL."""
    parsed = urlparse(url)
    if parsed.hostname in ("www.youtube.com", "youtube.com", "m.youtube.com"):
        return parse_qs(parsed.query).get("v", [None])[0]
    if parsed.hostname in ("youtu.be",):
        return parsed.path.lstrip("/")
    return None


def fetch_video_metadata(url):
    """Fetch video title, description, channel, duration, and thumbnail URLs."""
    ydl_opts = {"quiet": True, "skip_download": True, "no_warnings": True}
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=False)

    thumbnails = info.get("thumbnails", [])
    # Pick the best resolution thumbnails (sorted by preference/resolution)
    thumbnail_urls = [t["url"] for t in thumbnails if "url" in t]

    return {
        "title": info.get("title", "Untitled Video"),
        "description": info.get("description", ""),
        "channel": info.get("channel", info.get("uploader", "Unknown")),
        "duration": info.get("duration", 0),
        "thumbnail_urls": thumbnail_urls,
        "tags": info.get("tags", []),
        "categories": info.get("categories", []),
        "upload_date": info.get("upload_date", ""),
    }


def fetch_transcript(video_id):
    """Fetch the transcript for a YouTube video."""
    transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)

    # Prefer manually created transcripts, fall back to auto-generated
    try:
        transcript = transcript_list.find_manually_created_transcript(["en"])
    except Exception:
        try:
            transcript = transcript_list.find_generated_transcript(["en"])
        except Exception:
            # Try any available transcript
            for t in transcript_list:
                transcript = t
                break
            else:
                return None

    entries = transcript.fetch()
    return entries


def format_duration(seconds):
    """Format seconds into HH:MM:SS or MM:SS."""
    if seconds is None or seconds == 0:
        return "Unknown"
    hours, remainder = divmod(int(seconds), 3600)
    minutes, secs = divmod(remainder, 60)
    if hours > 0:
        return f"{hours}:{minutes:02d}:{secs:02d}"
    return f"{minutes}:{secs:02d}"


def build_full_text(transcript_entries):
    """Combine transcript entries into full text."""
    return " ".join(entry.get("text", "") for entry in transcript_entries)


def split_into_sentences(text):
    """Split text into sentences using basic punctuation rules."""
    # Split on sentence-ending punctuation followed by whitespace
    sentences = re.split(r"(?<=[.!?])\s+", text)
    return [s.strip() for s in sentences if s.strip()]


def extractive_summary(text, num_sentences=15):
    """Generate an extractive summary by scoring sentences on keyword frequency.

    This uses a simple TF-based approach: sentences containing the most frequent
    meaningful words are ranked higher and selected for the summary.
    """
    sentences = split_into_sentences(text)
    if len(sentences) <= num_sentences:
        return text

    # Compute word frequencies (simple TF)
    words = re.findall(r"\b[a-z]{3,}\b", text.lower())
    # Filter out very common English stop words
    stop_words = {
        "the", "and", "for", "are", "but", "not", "you", "all", "can", "her",
        "was", "one", "our", "out", "has", "have", "had", "this", "that",
        "with", "from", "they", "been", "said", "each", "which", "their",
        "will", "other", "about", "many", "then", "them", "these", "some",
        "would", "make", "like", "into", "could", "time", "very", "when",
        "what", "your", "just", "know", "take", "people", "come", "than",
        "look", "only", "its", "over", "think", "also", "back", "after",
        "use", "two", "how", "way", "who", "did", "get", "more", "going",
        "really", "thing", "things", "right", "there", "here", "where",
        "does", "because", "don", "well", "still", "should",
    }
    filtered = [w for w in words if w not in stop_words]
    freq = {}
    for w in filtered:
        freq[w] = freq.get(w, 0) + 1

    # Score each sentence
    scored = []
    for i, sentence in enumerate(sentences):
        s_words = re.findall(r"\b[a-z]{3,}\b", sentence.lower())
        score = sum(freq.get(w, 0) for w in s_words if w not in stop_words)
        # Slight bias toward earlier sentences (they often introduce topics)
        position_bonus = max(0, 1.0 - (i / len(sentences)) * 0.3)
        scored.append((i, sentence, score * position_bonus))

    # Select top sentences, maintain original order
    scored.sort(key=lambda x: x[2], reverse=True)
    selected = sorted(scored[:num_sentences], key=lambda x: x[0])

    return " ".join(s[1] for s in selected)


def segment_transcript(transcript_entries, segment_minutes=5):
    """Segment transcript into time-based chunks."""
    segments = []
    current_segment = None
    segment_duration = segment_minutes * 60

    for entry in transcript_entries:
        start = entry.get("start", 0)
        text = entry.get("text", "")

        segment_index = int(start // segment_duration)
        expected_start = segment_index * segment_duration

        if current_segment is None or expected_start != current_segment["start"]:
            if current_segment is not None:
                segments.append(current_segment)
            current_segment = {"start": expected_start, "texts": []}

        current_segment["texts"].append(text)

    if current_segment is not None and current_segment["texts"]:
        segments.append(current_segment)

    return segments


def download_image(url):
    """Download an image from a URL and return it as bytes."""
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        )
    }
    req = Request(url, headers=headers)
    with urlopen(req, timeout=15) as response:
        return response.read()


def get_best_thumbnails(thumbnail_urls, max_images=5):
    """Download the best available thumbnails, returning up to max_images."""
    images = []
    # Try thumbnails from best to worst (yt-dlp usually lists low-res first)
    for url in reversed(thumbnail_urls):
        if len(images) >= max_images:
            break
        try:
            data = download_image(url)
            images.append(data)
        except Exception:
            continue
    return images


def get_keyframe_thumbnails(video_id, count=4):
    """Get YouTube-hosted thumbnail images at different points in the video.

    YouTube auto-generates thumbnails at specific storyboard positions.
    We use well-known thumbnail URL patterns.
    """
    base = f"https://img.youtube.com/vi/{video_id}"
    patterns = [
        f"{base}/maxresdefault.jpg",
        f"{base}/sddefault.jpg",
        f"{base}/hqdefault.jpg",
        f"{base}/mqdefault.jpg",
        f"{base}/default.jpg",
        f"{base}/0.jpg",
        f"{base}/1.jpg",
        f"{base}/2.jpg",
        f"{base}/3.jpg",
    ]

    images = []
    for url in patterns:
        if len(images) >= count:
            break
        try:
            data = download_image(url)
            # Skip tiny placeholder images (< 1KB likely means error)
            if len(data) > 1024:
                images.append(data)
        except Exception:
            continue
    return images


def create_docx(metadata, summary, segments, images, output_path):
    """Create a .docx file with the video summary and images."""
    doc = Document()

    # --- Title ---
    title_para = doc.add_heading(metadata["title"], level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Video Info ---
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(
        f"Channel: {metadata['channel']}  |  "
        f"Duration: {format_duration(metadata['duration'])}"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    if metadata.get("upload_date"):
        date_str = metadata["upload_date"]
        if len(date_str) == 8:
            date_str = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"
        date_run = info_para.add_run(f"  |  Uploaded: {date_str}")
        date_run.font.size = Pt(11)
        date_run.font.color.rgb = RGBColor(100, 100, 100)

    if metadata.get("categories"):
        cat_para = doc.add_paragraph()
        cat_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cat_run = cat_para.add_run(
            f"Categories: {', '.join(metadata['categories'])}"
        )
        cat_run.font.size = Pt(10)
        cat_run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()  # spacer

    # --- Main Thumbnail ---
    if images:
        try:
            img_stream = BytesIO(images[0])
            doc.add_picture(img_stream, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    doc.add_paragraph()  # spacer

    # --- Executive Summary ---
    doc.add_heading("Executive Summary", level=1)
    for para_text in textwrap.wrap(summary, width=500):
        doc.add_paragraph(para_text)

    doc.add_paragraph()  # spacer

    # --- Additional Images ---
    if len(images) > 1:
        doc.add_heading("Video Snapshots", level=1)
        for i, img_data in enumerate(images[1:], start=1):
            try:
                img_stream = BytesIO(img_data)
                doc.add_picture(img_stream, width=Inches(4.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph(f"Figure {i}")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.runs[0] if cap.runs else cap.add_run(f"Figure {i}")
                cap_run.font.size = Pt(9)
                cap_run.font.color.rgb = RGBColor(130, 130, 130)
            except Exception:
                continue

    # --- Detailed Section Summaries ---
    if segments:
        doc.add_heading("Detailed Section Breakdown", level=1)
        for seg in segments:
            start_time = format_duration(seg["start"])
            section_text = " ".join(seg["texts"])
            section_summary = extractive_summary(section_text, num_sentences=5)

            doc.add_heading(f"Section starting at {start_time}", level=2)
            doc.add_paragraph(section_summary)

    # --- Video Description ---
    if metadata.get("description"):
        doc.add_heading("Original Video Description", level=1)
        # Limit description length for the document
        desc = metadata["description"]
        if len(desc) > 3000:
            desc = desc[:3000] + "..."
        doc.add_paragraph(desc)

    # --- Tags ---
    if metadata.get("tags"):
        doc.add_heading("Tags", level=2)
        doc.add_paragraph(", ".join(metadata["tags"][:30]))

    # Save
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Summarize a YouTube video and generate a .docx report."
    )
    parser.add_argument(
        "url",
        nargs="?",
        default=DEFAULT_VIDEO_URL,
        help="YouTube video URL (default: %(default)s)",
    )
    parser.add_argument(
        "-o", "--output",
        default=None,
        help="Output .docx file path (default: output/<video_title>.docx)",
    )
    parser.add_argument(
        "-s", "--sentences",
        type=int,
        default=15,
        help="Number of sentences in the executive summary (default: 15)",
    )
    args = parser.parse_args()

    video_url = args.url
    video_id = extract_video_id(video_url)
    if not video_id:
        print(f"Error: Could not extract video ID from URL: {video_url}")
        sys.exit(1)

    print(f"Video ID: {video_id}")
    print(f"Video URL: {video_url}")
    print()

    # Step 1: Fetch video metadata
    print("Fetching video metadata...")
    try:
        metadata = fetch_video_metadata(video_url)
        print(f"  Title: {metadata['title']}")
        print(f"  Channel: {metadata['channel']}")
        print(f"  Duration: {format_duration(metadata['duration'])}")
    except Exception as e:
        print(f"Error fetching video metadata: {e}")
        sys.exit(1)

    # Step 2: Fetch transcript
    print("\nFetching transcript...")
    try:
        transcript_entries = fetch_transcript(video_id)
        if not transcript_entries:
            print("Error: No transcript available for this video.")
            sys.exit(1)
        full_text = build_full_text(transcript_entries)
        print(f"  Transcript length: {len(full_text)} characters")
    except Exception as e:
        print(f"Error fetching transcript: {e}")
        sys.exit(1)

    # Step 3: Generate summary
    print("\nGenerating summary...")
    summary = extractive_summary(full_text, num_sentences=args.sentences)
    print(f"  Summary length: {len(summary)} characters")

    # Step 4: Segment transcript
    print("\nSegmenting transcript...")
    segments = segment_transcript(transcript_entries, segment_minutes=5)
    print(f"  Number of segments: {len(segments)}")

    # Step 5: Fetch images
    print("\nFetching images...")
    images = []

    # Try keyframe thumbnails first
    keyframes = get_keyframe_thumbnails(video_id, count=4)
    images.extend(keyframes)
    print(f"  Keyframe thumbnails: {len(keyframes)}")

    # Also try thumbnails from metadata
    if metadata.get("thumbnail_urls"):
        meta_imgs = get_best_thumbnails(metadata["thumbnail_urls"], max_images=2)
        images.extend(meta_imgs)
        print(f"  Metadata thumbnails: {len(meta_imgs)}")

    # Deduplicate (by size as a simple heuristic)
    seen_sizes = set()
    unique_images = []
    for img in images:
        if len(img) not in seen_sizes:
            seen_sizes.add(len(img))
            unique_images.append(img)
    images = unique_images
    print(f"  Total unique images: {len(images)}")

    # Step 6: Create output directory and docx
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    if args.output:
        output_path = args.output
    else:
        # Sanitize title for filename
        safe_title = re.sub(r'[^\w\s-]', '', metadata["title"])
        safe_title = re.sub(r'\s+', '_', safe_title).strip("_")[:80]
        output_path = os.path.join(OUTPUT_DIR, f"{safe_title}_summary.docx")

    print(f"\nCreating document: {output_path}")
    create_docx(metadata, summary, segments, images, output_path)
    print(f"\nDone! Document saved to: {output_path}")


if __name__ == "__main__":
    main()
