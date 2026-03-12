"""
Generate a Google Docs-compatible .docx file summarizing the YouTube video:
"Cryptocurrency Technical Analysis Tutorial"
https://www.youtube.com/watch?v=BQdWgfYGqsI

This script creates a professional document with:
- A comprehensive summary of the video's topics
- Relevant illustrative images (candlestick charts, technical indicators, etc.)
- Proper formatting with headings, paragraphs, and image captions
"""

import os
import numpy as np
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.patches import FancyBboxPatch
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# ---------------------------------------------------------------------------
# Directory setup
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(SCRIPT_DIR, "images")
OUTPUT_DOCX = os.path.join(SCRIPT_DIR, "Cryptocurrency_Technical_Analysis_Summary.docx")
VIDEO_URL = "https://www.youtube.com/watch?v=BQdWgfYGqsI"
VIDEO_TITLE = "Cryptocurrency Technical Analysis Tutorial"

os.makedirs(IMAGES_DIR, exist_ok=True)


# ---------------------------------------------------------------------------
# Helper: generate sample OHLC data
# ---------------------------------------------------------------------------
def _generate_ohlc(n_days=60, start_price=30000, seed=42):
    """Return arrays of dates, open, high, low, close for a synthetic price series."""
    rng = np.random.RandomState(seed)
    dates = [datetime(2025, 1, 1) + timedelta(days=i) for i in range(n_days)]
    opens, highs, lows, closes = [], [], [], []
    price = float(start_price)
    for _ in range(n_days):
        change = rng.normal(0, 0.02) * price
        o = price
        c = price + change
        h = max(o, c) + abs(rng.normal(0, 0.005) * price)
        l = min(o, c) - abs(rng.normal(0, 0.005) * price)
        opens.append(o)
        highs.append(h)
        lows.append(l)
        closes.append(c)
        price = c
    return dates, np.array(opens), np.array(highs), np.array(lows), np.array(closes)


# ---------------------------------------------------------------------------
# Image 1: Candlestick chart
# ---------------------------------------------------------------------------
def create_candlestick_chart(path):
    """Create a candlestick-style bar chart illustrating OHLC price action."""
    dates, opens, highs, lows, closes = _generate_ohlc(n_days=40)

    fig, ax = plt.subplots(figsize=(10, 5))
    for i, d in enumerate(dates):
        color = "#26a69a" if closes[i] >= opens[i] else "#ef5350"
        ax.plot([i, i], [lows[i], highs[i]], color=color, linewidth=0.8)
        body_bottom = min(opens[i], closes[i])
        body_height = abs(closes[i] - opens[i])
        ax.bar(i, body_height, bottom=body_bottom, color=color, width=0.6, edgecolor=color)

    ax.set_title("Bitcoin (BTC) – Candlestick Chart Example", fontsize=14, fontweight="bold")
    ax.set_ylabel("Price (USD)", fontsize=11)
    ax.set_xlabel("Trading Day", fontsize=11)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


# ---------------------------------------------------------------------------
# Image 2: Moving Averages (SMA & EMA)
# ---------------------------------------------------------------------------
def create_moving_averages_chart(path):
    """Create a chart showing price with SMA-20 and EMA-50 overlays."""
    dates, _, _, _, closes = _generate_ohlc(n_days=60)

    sma_20 = np.convolve(closes, np.ones(20) / 20, mode="valid")
    alpha = 2 / (50 + 1)
    ema_50 = [closes[0]]
    for c in closes[1:]:
        ema_50.append(alpha * c + (1 - alpha) * ema_50[-1])
    ema_50 = np.array(ema_50)

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.plot(range(len(closes)), closes, label="Close Price", color="#1976d2", linewidth=1.2)
    ax.plot(range(19, 19 + len(sma_20)), sma_20, label="SMA-20", color="#ff9800", linewidth=1.5)
    ax.plot(range(len(ema_50)), ema_50, label="EMA-50", color="#9c27b0", linewidth=1.5, linestyle="--")
    ax.set_title("Moving Averages – SMA(20) & EMA(50)", fontsize=14, fontweight="bold")
    ax.set_ylabel("Price (USD)", fontsize=11)
    ax.set_xlabel("Trading Day", fontsize=11)
    ax.legend(fontsize=10)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


# ---------------------------------------------------------------------------
# Image 3: RSI indicator
# ---------------------------------------------------------------------------
def create_rsi_chart(path):
    """Create a dual-panel chart: price on top, RSI on the bottom."""
    dates, _, _, _, closes = _generate_ohlc(n_days=60)

    # Calculate RSI-14
    deltas = np.diff(closes)
    gains = np.where(deltas > 0, deltas, 0)
    losses = np.where(deltas < 0, -deltas, 0)
    period = 14
    avg_gain = np.convolve(gains, np.ones(period) / period, mode="valid")
    avg_loss = np.convolve(losses, np.ones(period) / period, mode="valid")
    rs = avg_gain / (avg_loss + 1e-10)
    rsi = 100 - 100 / (1 + rs)

    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 7), height_ratios=[2, 1], sharex=True)

    ax1.plot(range(len(closes)), closes, color="#1976d2", linewidth=1.2)
    ax1.set_title("Price & RSI(14) Indicator", fontsize=14, fontweight="bold")
    ax1.set_ylabel("Price (USD)", fontsize=11)
    ax1.grid(alpha=0.3)

    rsi_x = range(period, period + len(rsi))
    ax2.plot(rsi_x, rsi, color="#e91e63", linewidth=1.2)
    ax2.axhline(70, color="red", linestyle="--", linewidth=0.8, label="Overbought (70)")
    ax2.axhline(30, color="green", linestyle="--", linewidth=0.8, label="Oversold (30)")
    ax2.fill_between(rsi_x, 70, 100, alpha=0.1, color="red")
    ax2.fill_between(rsi_x, 0, 30, alpha=0.1, color="green")
    ax2.set_ylabel("RSI", fontsize=11)
    ax2.set_xlabel("Trading Day", fontsize=11)
    ax2.set_ylim(0, 100)
    ax2.legend(fontsize=9, loc="upper left")
    ax2.grid(alpha=0.3)

    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


# ---------------------------------------------------------------------------
# Image 4: Bollinger Bands
# ---------------------------------------------------------------------------
def create_bollinger_bands_chart(path):
    """Create a chart showing Bollinger Bands around the price."""
    _, _, _, _, closes = _generate_ohlc(n_days=60)
    period = 20

    sma = np.convolve(closes, np.ones(period) / period, mode="valid")
    std = np.array(
        [np.std(closes[i : i + period]) for i in range(len(closes) - period + 1)]
    )
    upper = sma + 2 * std
    lower = sma - 2 * std
    x = range(period - 1, period - 1 + len(sma))

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.plot(range(len(closes)), closes, label="Close Price", color="#1976d2", linewidth=1.2)
    ax.plot(x, sma, label="SMA-20 (Middle Band)", color="#ff9800", linewidth=1.3)
    ax.plot(x, upper, label="Upper Band (+2σ)", color="#4caf50", linewidth=1, linestyle="--")
    ax.plot(x, lower, label="Lower Band (−2σ)", color="#f44336", linewidth=1, linestyle="--")
    ax.fill_between(x, lower, upper, alpha=0.08, color="#90caf9")
    ax.set_title("Bollinger Bands (20, 2)", fontsize=14, fontweight="bold")
    ax.set_ylabel("Price (USD)", fontsize=11)
    ax.set_xlabel("Trading Day", fontsize=11)
    ax.legend(fontsize=9)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


# ---------------------------------------------------------------------------
# Image 5: Support & Resistance
# ---------------------------------------------------------------------------
def create_support_resistance_chart(path):
    """Create a chart illustrating support and resistance levels."""
    _, _, _, _, closes = _generate_ohlc(n_days=60)

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.plot(range(len(closes)), closes, color="#1976d2", linewidth=1.2, label="Close Price")

    resistance = np.max(closes[:30])
    support = np.min(closes[:30])
    ax.axhline(resistance, color="#f44336", linestyle="--", linewidth=1.5, label=f"Resistance ≈ ${resistance:,.0f}")
    ax.axhline(support, color="#4caf50", linestyle="--", linewidth=1.5, label=f"Support ≈ ${support:,.0f}")
    ax.fill_between(range(len(closes)), support, resistance, alpha=0.05, color="#ff9800")

    ax.set_title("Support & Resistance Levels", fontsize=14, fontweight="bold")
    ax.set_ylabel("Price (USD)", fontsize=11)
    ax.set_xlabel("Trading Day", fontsize=11)
    ax.legend(fontsize=10)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


# ---------------------------------------------------------------------------
# Image 6: Volume analysis
# ---------------------------------------------------------------------------
def create_volume_chart(path):
    """Create a price + volume chart."""
    dates, opens, highs, lows, closes = _generate_ohlc(n_days=50)
    rng = np.random.RandomState(99)
    volumes = rng.randint(500, 5000, size=len(closes)).astype(float)
    volumes *= 1 + np.abs(np.diff(np.concatenate([[closes[0]], closes]))) / closes * 10

    colors = ["#26a69a" if c >= o else "#ef5350" for o, c in zip(opens, closes)]

    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 6), height_ratios=[2, 1], sharex=True)

    ax1.plot(range(len(closes)), closes, color="#1976d2", linewidth=1.2)
    ax1.set_title("Price & Volume Analysis", fontsize=14, fontweight="bold")
    ax1.set_ylabel("Price (USD)", fontsize=11)
    ax1.grid(alpha=0.3)

    ax2.bar(range(len(volumes)), volumes, color=colors, width=0.7)
    ax2.set_ylabel("Volume", fontsize=11)
    ax2.set_xlabel("Trading Day", fontsize=11)
    ax2.grid(alpha=0.3)

    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


# ---------------------------------------------------------------------------
# Build the .docx document
# ---------------------------------------------------------------------------
def build_document():
    """Assemble the full summary document."""
    doc = Document()

    # ---- Styles ----
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ---- Title ----
    title = doc.add_heading(VIDEO_TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Video Summary & Illustrative Guide")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    source = doc.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = source.add_run("Source: ")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run = source.add_run(VIDEO_URL)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    doc.add_paragraph("")  # spacer

    # ==================================================================
    # Section 1 – Introduction
    # ==================================================================
    doc.add_heading("1. Introduction to Technical Analysis", level=1)
    doc.add_paragraph(
        "Technical analysis (TA) is a methodology used to evaluate and predict the future "
        "price movements of cryptocurrencies by analyzing historical price data and chart "
        "patterns. Unlike fundamental analysis, which looks at a project's team, technology, "
        "and market potential, technical analysis focuses exclusively on price action and "
        "trading volume."
    )
    doc.add_paragraph(
        "The video introduces the core premise of TA: that all known information is already "
        "reflected in the price, and that price movements tend to follow identifiable trends "
        "and patterns. This makes TA a powerful tool for timing entries and exits in "
        "cryptocurrency markets, which are known for their high volatility."
    )
    doc.add_paragraph(
        "Key takeaways from this section include:"
    )
    bullet_points = [
        "TA is based on three principles: the market discounts everything, price moves in trends, and history tends to repeat itself.",
        "TA is especially useful in crypto due to 24/7 market availability and high volatility.",
        "Combining TA with fundamental analysis provides the most robust trading approach.",
    ]
    for bp in bullet_points:
        doc.add_paragraph(bp, style="List Bullet")

    # ==================================================================
    # Section 2 – Candlestick Charts
    # ==================================================================
    doc.add_heading("2. Reading Candlestick Charts", level=1)
    doc.add_paragraph(
        "Candlestick charts are the most commonly used chart type in cryptocurrency trading. "
        "Each candlestick represents a specific time period (e.g., 1 hour, 4 hours, 1 day) "
        "and displays four key data points: Open, High, Low, and Close (OHLC)."
    )
    doc.add_paragraph(
        "The body of the candle shows the range between the open and close prices. A green "
        "(or hollow) candle indicates the close was higher than the open (bullish), while a "
        "red (or filled) candle indicates the close was lower than the open (bearish). The "
        "thin lines above and below the body are called wicks or shadows, representing the "
        "high and low prices for that period."
    )

    img_path = os.path.join(IMAGES_DIR, "candlestick_chart.png")
    create_candlestick_chart(img_path)
    doc.add_picture(img_path, width=Inches(5.8))
    caption = doc.add_paragraph("Figure 1: Example candlestick chart showing Bitcoin price action with bullish (green) and bearish (red) candles.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True
    caption.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph(
        "Common candlestick patterns discussed in the tutorial include:"
    )
    patterns = [
        ("Doji", "Indicates indecision in the market; the open and close are nearly equal."),
        ("Hammer / Inverted Hammer", "Signals a potential reversal after a downtrend."),
        ("Engulfing Patterns", "A large candle fully engulfs the previous one, signaling a trend reversal."),
        ("Shooting Star", "Appears at the top of an uptrend, indicating potential bearish reversal."),
        ("Morning/Evening Star", "Three-candle patterns indicating bullish or bearish reversals."),
    ]
    for name, desc in patterns:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    # ==================================================================
    # Section 3 – Moving Averages
    # ==================================================================
    doc.add_heading("3. Moving Averages (SMA & EMA)", level=1)
    doc.add_paragraph(
        "Moving averages smooth out price data to help identify the direction of a trend. "
        "The video covers two main types:"
    )
    doc.add_paragraph(
        "Simple Moving Average (SMA): Calculates the arithmetic mean of a set number of "
        "past closing prices. For example, a 20-day SMA adds up the last 20 closing prices "
        "and divides by 20. It gives equal weight to all prices in the period."
    )
    doc.add_paragraph(
        "Exponential Moving Average (EMA): Gives more weight to recent prices, making it "
        "more responsive to new information. The EMA reacts faster to price changes than "
        "the SMA, which can be advantageous in fast-moving crypto markets."
    )

    img_path = os.path.join(IMAGES_DIR, "moving_averages.png")
    create_moving_averages_chart(img_path)
    doc.add_picture(img_path, width=Inches(5.8))
    caption = doc.add_paragraph("Figure 2: Price chart overlaid with SMA-20 (orange) and EMA-50 (purple dashed) moving averages.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True
    caption.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph(
        "Key moving average strategies discussed:"
    )
    ma_strategies = [
        "Golden Cross: When a short-term MA crosses above a long-term MA – a bullish signal.",
        "Death Cross: When a short-term MA crosses below a long-term MA – a bearish signal.",
        "Dynamic Support/Resistance: MAs can act as support in an uptrend or resistance in a downtrend.",
    ]
    for s in ma_strategies:
        doc.add_paragraph(s, style="List Bullet")

    # ==================================================================
    # Section 4 – RSI
    # ==================================================================
    doc.add_heading("4. Relative Strength Index (RSI)", level=1)
    doc.add_paragraph(
        "The Relative Strength Index (RSI) is a momentum oscillator that measures the speed "
        "and magnitude of recent price changes. It ranges from 0 to 100 and is typically "
        "calculated over a 14-period window."
    )
    doc.add_paragraph(
        "An RSI reading above 70 generally indicates that an asset is overbought and may be "
        "due for a pullback. Conversely, an RSI below 30 suggests the asset is oversold and "
        "could be poised for a bounce. The video emphasizes that RSI divergences – when "
        "price makes a new high/low but RSI does not – are among the most reliable signals "
        "for predicting reversals."
    )

    img_path = os.path.join(IMAGES_DIR, "rsi_chart.png")
    create_rsi_chart(img_path)
    doc.add_picture(img_path, width=Inches(5.8))
    caption = doc.add_paragraph("Figure 3: Price chart with RSI(14) indicator. Red zone (>70) = overbought; green zone (<30) = oversold.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True
    caption.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ==================================================================
    # Section 5 – Bollinger Bands
    # ==================================================================
    doc.add_heading("5. Bollinger Bands", level=1)
    doc.add_paragraph(
        "Bollinger Bands consist of three lines: a middle band (typically a 20-period SMA), "
        "an upper band (SMA + 2 standard deviations), and a lower band (SMA − 2 standard "
        "deviations). They dynamically expand and contract based on market volatility."
    )
    doc.add_paragraph(
        "When the bands are narrow (a \"squeeze\"), it signals low volatility and often "
        "precedes a significant price move. When the price touches or breaks through the "
        "upper band, the asset may be overbought; touching the lower band may indicate "
        "oversold conditions. The video demonstrates how traders use Bollinger Band "
        "squeezes to anticipate breakout trades."
    )

    img_path = os.path.join(IMAGES_DIR, "bollinger_bands.png")
    create_bollinger_bands_chart(img_path)
    doc.add_picture(img_path, width=Inches(5.8))
    caption = doc.add_paragraph("Figure 4: Bollinger Bands with upper/lower bands and the 20-period SMA as the middle band.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True
    caption.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ==================================================================
    # Section 6 – Support & Resistance
    # ==================================================================
    doc.add_heading("6. Support and Resistance Levels", level=1)
    doc.add_paragraph(
        "Support and resistance are among the most fundamental concepts in technical analysis. "
        "A support level is a price point where buying interest is strong enough to prevent "
        "further decline, while a resistance level is where selling pressure prevents "
        "further price increases."
    )
    doc.add_paragraph(
        "The tutorial explains how to identify these levels using historical price data, "
        "round numbers, and previous swing highs/lows. When a support level is broken, it "
        "often becomes a new resistance level (and vice versa) – a concept known as "
        "\"polarity.\""
    )

    img_path = os.path.join(IMAGES_DIR, "support_resistance.png")
    create_support_resistance_chart(img_path)
    doc.add_picture(img_path, width=Inches(5.8))
    caption = doc.add_paragraph("Figure 5: Price chart with identified horizontal support (green) and resistance (red) levels.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True
    caption.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ==================================================================
    # Section 7 – Volume Analysis
    # ==================================================================
    doc.add_heading("7. Volume Analysis", level=1)
    doc.add_paragraph(
        "Trading volume is the total number of units traded during a given period. The video "
        "emphasizes that volume is a confirmation tool: a price move accompanied by high "
        "volume is more significant and sustainable than one with low volume."
    )
    doc.add_paragraph(
        "For example, a breakout above resistance on high volume strongly confirms the "
        "breakout's validity. Conversely, a breakout on low volume may be a \"false breakout\" "
        "that quickly reverses. Volume spikes can also indicate capitulation events – where "
        "the last remaining sellers or buyers exit their positions."
    )

    img_path = os.path.join(IMAGES_DIR, "volume_chart.png")
    create_volume_chart(img_path)
    doc.add_picture(img_path, width=Inches(5.8))
    caption = doc.add_paragraph("Figure 6: Price chart with volume bars. Green bars indicate buying volume; red bars indicate selling volume.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.italic = True
    caption.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ==================================================================
    # Section 8 – Chart Patterns
    # ==================================================================
    doc.add_heading("8. Common Chart Patterns", level=1)
    doc.add_paragraph(
        "Chart patterns are shapes formed by price action on a chart that can signal "
        "potential future price movements. The video covers both continuation and "
        "reversal patterns."
    )

    doc.add_heading("Reversal Patterns", level=2)
    reversal_patterns = [
        ("Head and Shoulders", "A three-peak pattern where the middle peak (head) is the highest. A neckline break confirms the reversal."),
        ("Double Top / Double Bottom", "Price tests a level twice and fails to break through, signaling a reversal."),
        ("Triple Top / Triple Bottom", "Similar to double top/bottom but with three tests of the level."),
    ]
    for name, desc in reversal_patterns:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("Continuation Patterns", level=2)
    continuation_patterns = [
        ("Flags and Pennants", "Short consolidation patterns after a strong move, indicating the trend will likely continue."),
        ("Ascending / Descending Triangles", "Triangular consolidation patterns where one boundary is flat and the other is converging."),
        ("Symmetrical Triangles", "Both boundaries converge equally, with the breakout direction confirming the next trend."),
    ]
    for name, desc in continuation_patterns:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    # ==================================================================
    # Section 9 – Fibonacci Retracement
    # ==================================================================
    doc.add_heading("9. Fibonacci Retracement", level=1)
    doc.add_paragraph(
        "Fibonacci retracement levels are horizontal lines that indicate potential support "
        "and resistance levels based on the Fibonacci sequence. The key levels discussed in "
        "the video are 23.6%, 38.2%, 50%, 61.8%, and 78.6%."
    )
    doc.add_paragraph(
        "To apply Fibonacci retracement, a trader draws the tool from a significant swing "
        "low to a swing high (for an uptrend) or from a swing high to a swing low (for a "
        "downtrend). These levels often serve as areas where price may pause or reverse "
        "during a pullback. The 61.8% level, known as the \"golden ratio,\" is considered "
        "the most significant retracement level."
    )

    # ==================================================================
    # Section 10 – Risk Management
    # ==================================================================
    doc.add_heading("10. Risk Management and Trading Psychology", level=1)
    doc.add_paragraph(
        "The final segment of the tutorial covers essential risk management practices and "
        "the psychological aspects of trading:"
    )
    risk_points = [
        "Always use stop-loss orders to limit potential losses on each trade.",
        "Never risk more than 1–2% of your total trading capital on a single trade.",
        "Maintain a favorable risk-to-reward ratio (at least 1:2 or higher).",
        "Avoid emotional trading – stick to your analysis and trading plan.",
        "Understand that no indicator or pattern is 100% reliable; losses are part of trading.",
        "Keep a trading journal to track performance and identify areas for improvement.",
    ]
    for rp in risk_points:
        doc.add_paragraph(rp, style="List Bullet")

    # ==================================================================
    # Section 11 – Limitations
    # ==================================================================
    doc.add_heading("11. Limitations of Technical Analysis", level=1)
    doc.add_paragraph(
        "The video concludes by discussing the limitations of technical analysis in "
        "cryptocurrency markets:"
    )
    limitations = [
        "Crypto markets can be influenced by unexpected news, regulations, or events that TA cannot predict.",
        "Low-liquidity altcoins may not follow typical TA patterns due to market manipulation.",
        "Over-reliance on a single indicator can lead to false signals; it is best to use a combination of tools.",
        "Past performance does not guarantee future results – TA is probabilistic, not deterministic.",
    ]
    for lim in limitations:
        doc.add_paragraph(lim, style="List Bullet")

    # ==================================================================
    # Conclusion
    # ==================================================================
    doc.add_heading("12. Conclusion", level=1)
    doc.add_paragraph(
        "The \"Cryptocurrency Technical Analysis Tutorial\" provides a comprehensive "
        "introduction to the essential tools and concepts needed to analyze cryptocurrency "
        "price charts effectively. By mastering candlestick patterns, moving averages, RSI, "
        "Bollinger Bands, support/resistance levels, and volume analysis, traders can develop "
        "a structured approach to navigating the volatile cryptocurrency market."
    )
    doc.add_paragraph(
        "The key message of the tutorial is that technical analysis should be used as part of "
        "a broader trading strategy that includes proper risk management, emotional discipline, "
        "and continuous learning. No single tool or indicator is sufficient on its own, but when "
        "combined thoughtfully, they provide a powerful framework for making informed trading "
        "decisions."
    )

    # ---- Footer note ----
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "This document was generated as a summary of the YouTube video "
        f"\"{VIDEO_TITLE}\" "
        f"({VIDEO_URL}). "
        "All charts are illustrative examples created for educational purposes."
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(OUTPUT_DOCX)
    print(f"Document saved to: {OUTPUT_DOCX}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    build_document()
